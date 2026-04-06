import docx
import re
import os

def clean_word(raw_word):
    if not raw_word:
        return None
    raw_word = str(raw_word).replace('\n', ' ').strip()
    
    # 去除词性标记
    tags = [r'\badj\.', r'\badv\.', r'\bn\.', r'\bv\.', r'\bprep\.', r'\bconj\.', r'\bpron\.', r'\bnum\.', r'\bart\.', r'\bint\.', r'\bphr\.']
    for tag in tags:
        raw_word = re.sub(tag, '', raw_word)
    
    # 彻底去除星号和其他非字母修饰
    raw_word = raw_word.replace('*', '').strip()
    
    # 直接按第一个括号、中文或特殊符号切断
    clean_w = re.split(r'[\(（\[【\u4e00-\u9fa5]', raw_word)[0].strip()
    
    # 只保留纯字母的单词 (a-z, A-Z, 连字符)
    clean_w = re.sub(r'[^a-zA-Z\-]', '', clean_w)
    
    if not clean_w or len(clean_w) < 2:
        return None
    
    return clean_w.lower()

def extract_from_doc(input_path, output_path):
    doc = docx.Document(input_path)
    words = []
    
    # 遍历所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 尝试按空格/逗号/分号分割
            parts = re.split(r'[,，;；\s\n]+', text)
            for part in parts:
                cleaned = clean_word(part)
                if cleaned:
                    words.append(cleaned)
    
    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts = re.split(r'[,，;；\s\n]+', text)
                    for part in parts:
                        cleaned = clean_word(part)
                        if cleaned:
                            words.append(cleaned)
    
    # 全局去重，并按照字母顺序排列
    unique_words = sorted(list(set(words)))
    
    with open(output_path, 'w', encoding='utf-8') as f:
        for w in unique_words:
            f.write(w + "\n")
    
    print(f"✅ DOC 提取完成！")
    print(f"共发现 {len(words)} 条记录，去重后剩余 {len(unique_words)} 个单词。")
    print(f"最终结果已保存至: {output_path}")

if __name__ == '__main__':
    extract_from_doc('/Volumes/ssd/ppdc/tools/book/小学/译林版/最新译林版小学英语三-六年级单词汇总.doc', '/Volumes/ssd/ppdc/tools/book/pdf_pipeline/yilin_raw.txt')
